from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.oxml.ns as nsmap

FONT = "Tahoma"
FONT_SIZE = Pt(10)
LINE_SPACING = 1.15


def set_paragraph_format(para, bold=False, size=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = LINE_SPACING
    for run in para.runs:
        run.font.name = FONT
        run.font.size = size or FONT_SIZE
        run.font.bold = bold


def add_paragraph(doc, text, bold=False, size=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = size or FONT_SIZE
    run.font.bold = bold
    para.alignment = align
    para.paragraph_format.line_spacing = LINE_SPACING
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def add_heading(doc, text, level_size=Pt(12), space_before=10, space_after=4):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = level_size
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = LINE_SPACING
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def add_sub_heading(doc, text, space_before=6, space_after=2):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = FONT_SIZE
    run.font.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = LINE_SPACING
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = FONT_SIZE
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = LINE_SPACING
    para.paragraph_format.space_after = Pt(3)
    return para


doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─────────────────────────────────────────────────────────────────────────────
# TITLE BLOCK
# ─────────────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title_run = title.add_run("E-Commerce Inventory API")
title_run.font.name = FONT
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
sub_run = sub.add_run("Design Report — PROG315 Object-Oriented Programming 2")
sub_run.font.name = FONT
sub_run.font.size = Pt(10)
sub_run.font.bold = False
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)

info = doc.add_paragraph()
info_run = info.add_run(
    "Group D  |  Limkokwing University of Creative Technology, Sierra Leone  |  Semester 4, March–July 2026"
)
info_run.font.name = FONT
info_run.font.size = Pt(9)
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.paragraph_format.space_after = Pt(10)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — GROUP MEMBERS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1. Group Members")
table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Name"
hdr[1].text = "Student ID"
hdr[2].text = "Module / Role"
data = [
    ("Mohamed Barry",              "BIT1201XXX", "Project Lead – Auth, Database & Core Architecture"),
    ("Jaria Bah",                  "BIT1201XXX", "Products & Categories Module"),
    ("Abdul Hakeem Gibril Kargbo", "BIT1201XXX", "Orders Module & Documentation"),
]
for i, (name, sid, role) in enumerate(data, start=1):
    row = table.rows[i].cells
    row[0].text = name
    row[1].text = sid
    row[2].text = role

for row in table.rows:
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = FONT
                run.font.size = FONT_SIZE
            para.paragraph_format.line_spacing = LINE_SPACING

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — PROJECT OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2. Project Overview")
add_paragraph(doc,
    "This project is a FastAPI-based RESTful API developed as part of the PROG315 Object-Oriented "
    "Programming 2 module. The application is an E-Commerce Inventory API that enables businesses "
    "to manage products, categories, orders, and users through a secure and well-documented REST "
    "interface. The system targets local Sierra Leonean traders and small businesses, helping them "
    "digitise their operations and access wider markets, directly supporting SDG 8 – Decent Work "
    "and Economic Growth."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — DESIGN CHOICES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3. Design Choices")

add_sub_heading(doc, "3.1  FastAPI as the Core Framework")
add_paragraph(doc,
    "FastAPI was chosen for its high performance, native support for async/await, and automatic "
    "generation of interactive documentation via Swagger UI and ReDoc. Its integration with Pydantic "
    "for data validation and type checking significantly reduces boilerplate code and prevents "
    "invalid data from reaching the database layer."
)

add_sub_heading(doc, "3.2  PostgreSQL with SQLAlchemy ORM")
add_paragraph(doc,
    "PostgreSQL was selected as the relational database due to its robustness, support for complex "
    "queries, and wide production adoption. SQLAlchemy ORM was used to map Python classes directly "
    "to database tables, providing an abstraction layer that makes the codebase more maintainable "
    "and portable. The ORM also handles relationships — such as the one-to-many relationship "
    "between Category and Product, and between User and Order — declaratively using "
    "relationship() and ForeignKey columns."
)

add_sub_heading(doc, "3.3  Modular Router Structure")
add_paragraph(doc,
    "The application was split into four dedicated routers (users, categories, products, orders), "
    "each in its own file under the routers/ directory. This separation of concerns ensures that "
    "each team member could independently develop and test their module without causing merge "
    "conflicts. Each router is mounted with its own prefix (e.g., /api/v1/products) following "
    "RESTful naming conventions."
)

add_sub_heading(doc, "3.4  Authentication and Role-Based Access Control")
add_paragraph(doc,
    "Authentication is implemented using OAuth2 with JWT Bearer tokens. User passwords are hashed "
    "using bcrypt via the Passlib library before storage, ensuring that plaintext credentials are "
    "never persisted. Two roles were defined — admin and customer. Admin users have full CRUD "
    "access to all resources, while customer users can browse products, place orders, and view "
    "only their own orders. This role-based access is enforced through dependency injection using "
    "require_admin() and get_current_active_user() functions passed via Depends()."
)

add_sub_heading(doc, "3.5  Dependency Injection")
add_paragraph(doc,
    "FastAPI's built-in Depends() mechanism is used consistently across all endpoints to inject "
    "the database session (get_db) and the authenticated current user. This ensures that database "
    "connections are properly opened and closed per request, preventing connection leaks, and that "
    "authentication checks are enforced at the route level rather than inside business logic."
)

add_sub_heading(doc, "3.6  Asynchronous Programming")
add_paragraph(doc,
    "The product creation endpoint is implemented as an async function, demonstrating the use of "
    "async/await for I/O-bound tasks. A helper function fetch_external_price_check() simulates "
    "a non-blocking call to an external pricing service, illustrating how the API can be extended "
    "to integrate with third-party services without blocking the event loop."
)

add_sub_heading(doc, "3.7  Type Annotations and Pydantic Schemas")
add_paragraph(doc,
    "Python type hints are used throughout the application — in route function signatures, model "
    "definitions, and schema classes. Pydantic v2 schemas enforce input validation, such as "
    "minimum length constraints on strings, positive price values, and valid email addresses, "
    "returning structured error messages automatically when constraints are violated."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — SDG ALIGNMENT
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4. SDG Alignment — SDG 8: Decent Work and Economic Growth")
add_paragraph(doc,
    "The E-Commerce Inventory API directly supports United Nations Sustainable Development Goal 8 "
    "by providing a digital tool that empowers small and medium-sized enterprises (SMEs) and "
    "informal traders in Sierra Leone to manage their inventory, process customer orders, and "
    "monitor stock levels online. Many small businesses in Sierra Leone operate without any "
    "digital inventory system, leading to stockouts, overselling, and revenue loss. This API "
    "provides a free, open-source foundation that developers and entrepreneurs can adapt to build "
    "affordable e-commerce solutions tailored to the local market."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — CHALLENGES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5. Challenges and How We Overcame Them")

add_sub_heading(doc, "5.1  Database Relationship Mapping")
add_paragraph(doc,
    "Mapping the many-to-many relationship between orders and products required an intermediate "
    "OrderItem model. Understanding how SQLAlchemy handles cascade deletes and how to correctly "
    "set back_populates on both sides of a relationship took significant testing. We resolved this "
    "by carefully reading the SQLAlchemy documentation and writing isolated test cases to verify "
    "the behaviour before integrating it into the main application."
)

add_sub_heading(doc, "5.2  JWT Token Validation and OAuth2 Flow")
add_paragraph(doc,
    "Implementing the OAuth2PasswordRequestForm correctly so that Swagger UI's Authorize button "
    "would work required understanding that FastAPI expects the login endpoint to accept "
    "form-encoded data, not JSON. This was a non-obvious distinction that caused initial "
    "authentication failures. The issue was resolved by using OAuth2PasswordRequestForm as the "
    "dependency rather than a Pydantic model."
)

add_sub_heading(doc, "5.3  Stock Management During Order Creation")
add_paragraph(doc,
    "Ensuring that stock quantities are atomically deducted when an order is placed — and that "
    "orders with insufficient stock are rejected before any database writes occur — required "
    "careful sequencing of validation checks within the create_order endpoint. We solved this by "
    "validating all items in a loop before committing any changes, so partial stock deductions "
    "cannot occur."
)

add_sub_heading(doc, "5.4  Team Collaboration via GitHub")
add_paragraph(doc,
    "Coordinating three team members working on separate modules required clear branch naming "
    "conventions and pull request reviews to avoid merge conflicts. We assigned each member a "
    "dedicated module (auth/database, products/categories, orders/documentation) and used GitHub "
    "Issues to track progress before merging into the main branch."
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6. Conclusion")
add_paragraph(doc,
    "The E-Commerce Inventory API successfully meets all core requirements outlined in the PROG315 "
    "project brief. The application demonstrates FastAPI best practices including dependency "
    "injection, OAuth2/JWT authentication, SQLAlchemy ORM integration, async programming, Pydantic "
    "data validation, and full Swagger UI and ReDoc documentation. The project is published under "
    "the MIT open-source license and developed collaboratively on GitHub with individual "
    "contributions tracked through branches and pull requests. The system is aligned with SDG 8 "
    "and addresses a real digital gap faced by small businesses in Sierra Leone."
)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
doc.save("Group_D_PROG315_Report.docx")
print("Report saved: Group_D_PROG315_Report.docx")
